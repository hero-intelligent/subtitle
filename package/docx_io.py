from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from tqdm import tqdm
from io import BytesIO

from package.accept_review import accept_changes
from package.time_reform import time_str_to_ms, time_ms_to_str

import pandas as pd

def parse_docx_file(docx_file_path: str) -> pd.DataFrame:
    events_left = []
    events_right = []

    stream = BytesIO()
    accept_changes(docx_file_path, stream)
    stream.seek(0)

    doc = Document(stream)

    # 遍历表格中的每一行
    for row in tqdm(doc.tables[0].rows, desc="parsing docx file"):

        index = int(row.cells[0].text.strip())

        time = row.cells[1].text.split("-->",1).strip()

        start = time[0].strip()
        start = time_str_to_ms(start)
        end = time[1].strip()
        end = time_str_to_ms(end)

        text_left = row.cells[2].text.strip()
        text_right = row.cells[3].text.strip()

        event_left = {
            "Index": index,
            "Start": start,
            "End": end,
            "Text": text_left
        }

        event_right = {
            "Index": index,
            "Start": start,
            "End": end,
            "Text": text_right
        }

        events_left.append(event_left)
        events_right.append(event_right)

    return pd.DataFrame(events_left), pd.DataFrame(events_right)




# Function to convert JSON to Word table format
def json_to_word(
        original_events: list,
        translated_events: list = None,
        reference_data: dict = None,
        force_disgard_index_and_time: bool = True
        ) -> Document:
    

    doc = Document()

    # Create a table with 4 columns: Number, Time (start -> end), Subtitle Text, Empty
    table = doc.add_table(rows=len(original_events), cols=4 if reference_data is None else 5)
    table.style = 'Table Grid'

    # Loop through events and fill in the table
    for i, event in tqdm(enumerate(original_events), desc="filling original text in docx file", total=len(original_events)):
        index = event['Index']
        start_time = event['Start']
        end_time = event['End']
        subtitle_text = event['Text']

        row_cells = table.rows[i].cells

        row_cells[0].text = str(index)  # Number
        row_cells[1].text = f"{time_ms_to_str(start_time)} --> {time_ms_to_str(end_time)}"  # Time range
        row_cells[2].text = subtitle_text  # Subtitle text
        row_cells[3].text = ""  # Empty column

    if translated_events is not None:
        # Loop through events and fill in the table
        for i, event in tqdm(enumerate(translated_events), desc="filling translated text in docx file", total=len(translated_events)):
            index = event['Index']
            start_time = event['Start']
            end_time = event['End']
            subtitle_text = event['Text']

            row_cells = table.rows[i].cells

            is_index_same = row_cells[0].text == str(index)  # Number
            is_timestamp_same = row_cells[1].text == f"{time_ms_to_str(start_time)} --> {time_ms_to_str(end_time)}"  # Time range

            if force_disgard_index_and_time:
                row_cells[3].text = subtitle_text  # Empty column
            elif is_index_same and is_timestamp_same:
                row_cells[3].text = subtitle_text  # Empty column
            else:
                print("index or time match failed.")
                print("original data:")
                print(row_cells[0].text, row_cells[1].text, row_cells[2].text)
                print("translated data:")
                print(str(index), f"{time_ms_to_str(start_time)} --> {time_ms_to_str(end_time)}", subtitle_text)

    return doc

