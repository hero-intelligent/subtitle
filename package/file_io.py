import os
import copy
import json
import re
from typing import Union
import warnings
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from tqdm import tqdm

from zipfile import ZipFile
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
from docx import Document

from io import BytesIO

from package.accept_review import accept_changes

data_template = {
    "ScriptInfo": {},
    "Styles": [],
    "Events": []
}

def common_template() -> dict:
    default_script_info = {
        "ScriptType": "v4.00+",
        "WrapStyle": "0",
        "ScaledBorderAndShadow": "no",
        "Timer": "100.0000",
        "PlayResX": "1920",
        "PlayResY": "1080",
        "YCbCr Matrix": "None"
    }

    default_style = [
        {
            "Name": "Default",
            "Fontname": "Arial",
            "Fontsize": 20.0,
            "PrimaryColour": "&H00FFFFFF",
            "SecondaryColour": "&H0300FFFF",
            "OutlineColour": "&H00000000",
            "BackColour": "&H02000000",
            "Bold": 0,
            "Italic": 0,
            "Underline": 0,
            "StrikeOut": 0,
            "ScaleX": 100,
            "ScaleY": 100,
            "Spacing": 0,
            "Angle": 0,
            "BorderStyle": 1,
            "Outline": 2,
            "Shadow": 1,
            "Alignment": 2,
            "MarginL": 10,
            "MarginR": 10,
            "MarginV": 10,
            "Encoding": 1
        },
        {
            "Name": "Descriptive-920",
            "Fontname": "Noto Sans SC Medium",
            "Fontsize": 64.0,
            "PrimaryColour": "&H21FFFFFF",
            "SecondaryColour": "&H21FFFFFF",
            "OutlineColour": "&H00000000",
            "BackColour": "&H00000000",
            "Bold": 0,
            "Italic": 0,
            "Underline": 0,
            "StrikeOut": 0,
            "ScaleX": 100,
            "ScaleY": 100,
            "Spacing": 0,
            "Angle": 0,
            "BorderStyle": 1,
            "Outline": 0,
            "Shadow": 0,
            "Alignment": 2,
            "MarginL": 10,
            "MarginR": 10,
            "MarginV": 920,
            "Encoding": 1
        }
    ]

    data = copy.deepcopy(data_template)
    data["ScriptInfo"] = default_script_info
    data["Styles"] = default_style

    return copy.deepcopy(data)

event_template = {
    "Layer": 0,
    "Style": "Default",
    "Name": "",
    "MarginL": 0,
    "MarginR": 0,
    "MarginV": 0,
    "Effect": "",
}

def timestamp_reform(time_str: str, output_format: str = "srt") -> str:
    pattern_srt = r"^(\d{2}):(\d{2}):(\d{2}),(\d{3})$"
    pattern_ass = r"^(\d{1,2}):(\d{2}):(\d{2})\.(\d{2})$"
    pattern_wrong = r"^(\d+):(\d{2}):(\d{2})[:.,](\d+)$"

    # 使用正则表达式进行匹配
    match_srt = re.match(pattern_srt, time_str)
    match_ass = re.match(pattern_ass, time_str)
    match_wrong = re.match(pattern_wrong, time_str)

    if match_srt:
        # print("匹配srt成功！")
        hours = int(match_srt.group(1))
        minutes = int(match_srt.group(2))
        seconds = int(match_srt.group(3))
        milliseconds = int(round(float("0." + match_srt.group(4)) * 1000))
        # print(f"小时: {hours}, 分钟: {minutes}, 秒: {seconds}, 毫秒: {milliseconds}")
    elif match_ass:
        # print("匹配ass成功！")
        hours = int(match_ass.group(1))
        minutes = int(match_ass.group(2))
        seconds = int(match_ass.group(3))
        milliseconds = int(round(float("0." + match_ass.group(4)) * 1000))
        # print(f"小时: {hours}, 分钟: {minutes}, 秒: {seconds}, 毫秒: {milliseconds}")
    elif match_wrong:
        warnings.warn(f"Invalid format detected: {time_str}", UserWarning)
        hours = int(match_wrong.group(1))
        minutes = int(match_wrong.group(2))
        seconds = int(match_wrong.group(3))
        milliseconds = int(round(float("0." + match_wrong.group(4)) * 1000))
        print(f"小时: {hours}, 分钟: {minutes}, 秒: {seconds}, 毫秒: {milliseconds}")
    else:
        print(f"{time_str}匹配失败")
        return None

    if output_format == "srt":
        return f"{int(hours):02}:{int(minutes):02}:{int(seconds):02},{milliseconds:03}"
    elif output_format == "ass":
        return f"{int(hours):01}:{int(minutes):02}:{int(seconds):02}.{int(round(milliseconds / 1000, 2) * 100)}"

def parse_ass_file(ass_file_path: str) -> dict:
    data = copy.deepcopy(data_template)

    with open(ass_file_path, 'r', encoding='utf-8') as file:
        content = file.read()
        content = content.lstrip('\ufeff')  # Remove BOM if it's at the start

    section = None

    event_index = 0

    for line in tqdm(content.splitlines(), desc="parsing ass file"):
        line = line.strip()

        if not line or line.startswith(";"):  # Skip empty lines and comments
            continue

        if line.startswith("[Script Info]"):
            section = "ScriptInfo"
        elif line.startswith("[V4+ Styles]"):
            section = "Styles"
        elif line.startswith("[Events]"):
            section = "Events"

        # Parse Script Info
        if section == "ScriptInfo" and not line.startswith("["):
            if ":" in line:
                key, value = line.split(":", 1)
                data["ScriptInfo"][key.strip()] = value.strip()

        # Parse Styles
        elif section == "Styles" and line.startswith("Style:"):
            style_data = line.split(":", 1)[1].split(",", 24)  # Split into 24 elements (to match the style format)
            style = {
                "Name": style_data[0].strip(),
                "Fontname": style_data[1].strip(),
                "Fontsize": float(style_data[2].strip()),
                "PrimaryColour": style_data[3].strip(),
                "SecondaryColour": style_data[4].strip(),
                "OutlineColour": style_data[5].strip(),
                "BackColour": style_data[6].strip(),
                "Bold": int(style_data[7].strip()),
                "Italic": int(style_data[8].strip()),
                "Underline": int(style_data[9].strip()),
                "StrikeOut": int(style_data[10].strip()),
                "ScaleX": int(style_data[11].strip()),
                "ScaleY": int(style_data[12].strip()),
                "Spacing": int(style_data[13].strip()),
                "Angle": int(style_data[14].strip()),
                "BorderStyle": int(style_data[15].strip()),
                "Outline": int(style_data[16].strip()),
                "Shadow": int(style_data[17].strip()),
                "Alignment": int(style_data[18].strip()),
                "MarginL": int(style_data[19].strip()),
                "MarginR": int(style_data[20].strip()),
                "MarginV": int(style_data[21].strip()),
                "Encoding": int(style_data[22].strip())
            }
            data["Styles"].append(style)

        # Parse Events
        elif section == "Events" and line.startswith("Dialogue:"):
            event_index = event_index + 1
            event_data = line.split(":", 1)[1].split(",", 9)  # Split the dialogue data into 9 parts
            event = {
                "Index": event_index,
                "Layer": int(event_data[0].strip()),
                "Start": timestamp_reform(event_data[1].strip()),   # SRT style
                "End": timestamp_reform(event_data[2].strip()),     # SRT style
                "Style": event_data[3].strip(),
                "Name": event_data[4].strip(),
                "MarginL": int(event_data[5].strip()),
                "MarginR": int(event_data[6].strip()),
                "MarginV": int(event_data[7].strip()),
                "Effect": event_data[8].strip(),
                "Text": event_data[9].strip().replace("\\N","\n")       # SRT Style
            }
            data["Events"].append(event)

    return copy.deepcopy(data)

def parse_srt_file(srt_file_path: str) -> dict:
    data = copy.deepcopy(common_template())

    with open(srt_file_path, 'r', encoding='utf-8') as file:
        content = file.read()
        content = content.lstrip('\ufeff')  # Remove BOM if it's at the start
        content = content.strip()

    for line in tqdm(content.split("\n\n"), desc="parsing srt file"):
        line = line.strip()

        event_data = line.split("\n", 2)  # Split the suptitle into 3 parts

        index = event_data[0]
        start = event_data[1].split("-->",1)[0].strip()
        end = event_data[1].split("-->",1)[1].strip()

        if len(event_data) == 3:
            text = event_data[2]
        else:
            text = ""
            warnings.warn(f"line {index} is empty!", UserWarning)

        for c in "[]()（）【】":
            if c in text:
                style = "Descriptive-920"
            else:
                style = event_template.get("Style", "")

        event = {
            **event_template,
            "Index": index,
            "Start": start,
            "End": end,
            "Text": text,
            "Style": style
        }
        data["Events"].append(event)

    return copy.deepcopy(data)

def parse_docx_file(docx_file_path: str, both: bool = False, translated: bool = True) -> tuple[dict, dict]:
    data_left = copy.deepcopy(common_template())
    data_right = copy.deepcopy(common_template())

    stream = BytesIO()
    accept_changes(docx_file_path, stream)
    stream.seek(0)

    doc = Document(stream)

    rows_translated = 0
    rows_not_translated = 0

    # 遍历表格中的每一行
    for row in tqdm(doc.tables[0].rows, desc="parsing docx file"):

        if row.cells[3].text.strip():
            rows_translated +=1
        else:
            rows_not_translated += 1

        event_left = {
            **copy.deepcopy(event_template),
            "Index": int(row.cells[0].text.strip()),
            "Start": row.cells[1].text.split("-->",1)[0].strip(),
            "End": row.cells[1].text.split("-->",1)[1].strip(),
            "Text": row.cells[2].text.strip(),
            "Style": "Descriptive-920" if any(c in row.cells[2].text for c in "[]()（）【】") else event_template.get("Style", "")
        }

        event_right = {
            **copy.deepcopy(event_template),
            "Index": int(row.cells[0].text.strip()),
            "Start": row.cells[1].text.split("-->",1)[0].strip(),
            "End": row.cells[1].text.split("-->",1)[1].strip(),
            "Text": row.cells[3].text.strip(),
            "Style": "Descriptive-920" if any(c in row.cells[3].text for c in "[]()（）【】") else event_template.get("Style", "")
        }

        data_left["Events"].append(event_left)
        data_right["Events"].append(event_right)

    if rows_translated == 0:
        translated = False
    elif rows_not_translated / rows_translated >= 1:
        both = False

    return copy.deepcopy(data_left), copy.deepcopy(data_right)

    # if both:
    #     return [copy.deepcopy(data_left),copy.deepcopy(data_right)]
    # elif translated:
    #     return copy.deepcopy(data_right)
    # else:
    #     return copy.deepcopy(data_left)

def json_to_ass(json_data: dict) -> str:
    ass_output = []

    # Write Script Info
    script_info = json_data.get('ScriptInfo', {})

    ass_output.append("[Script Info]")
    for key, value in tqdm(script_info.items(), desc="generating script info in ass file"):
        ass_output.append(f"{key}: {value}")
    ass_output.append("")

    # Write Styles
    ass_output.append("[V4+ Styles]")
    ass_output.append("Format: Name,Fontname,Fontsize,PrimaryColour,SecondaryColour,OutlineColour,BackColour,Bold,Italic,Underline,StrikeOut,ScaleX,ScaleY,Spacing,Angle,BorderStyle,Outline,Shadow,Alignment,MarginL,MarginR,MarginV,Encoding")
    styles = json_data.get('Styles', [])
    for style in tqdm(styles, desc="generating style in ass file"):
        style_line = ",".join([
            str(style.get("Name", "")).strip(),
            str(style.get("Fontname", "")).strip(),
            str(style.get("Fontsize", "")).strip(),
            str(style.get("PrimaryColour", "")).strip(),
            str(style.get("SecondaryColour", "")).strip(),
            str(style.get("OutlineColour","")).strip(),
            str(style.get("BackColour", "")).strip(),
            str(style.get("Bold", "")).strip(),
            str(style.get("Italic", "")).strip(),
            str(style.get("Underline", "")).strip(),
            str(style.get("StrikeOut", "")).strip(),
            str(style.get("ScaleX", "")).strip(),
            str(style.get("ScaleY", "")).strip(),
            str(style.get("Spacing", "")).strip(),
            str(style.get("Angle", "")).strip(),
            str(style.get("BorderStyle", "")).strip(),
            str(style.get("Outline", "")).strip(),
            str(style.get("Shadow", "")).strip(),
            str(style.get("Alignment", "")).strip(),
            str(style.get("MarginL", "")).strip(),
            str(style.get("MarginR", "")).strip(),
            str(style.get("MarginV", "")).strip(),
            str(style.get("Encoding", "")).strip()
        ])
        ass_output.append(f"Style: {style_line}")
    ass_output.append("")

    # Write Events
    ass_output.append("[Events]")
    ass_output.append("Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text")
    events = json_data.get('Events', [])
    for event in tqdm(events, desc="generating event in ass file"):
        event_line = ",".join([
            str(event.get("Layer", "")).strip(),
            timestamp_reform(str(event.get("Start", "")).strip(), output_format="ass"),  # SRT style stored, ASS wanted
            timestamp_reform(str(event.get("End", "")).strip(), output_format="ass"),    # SRT style stored, ASS wanted
            str(event.get("Style", "")).strip(),
            str(event.get("Name", "")).strip(),
            str(event.get("MarginL", "")).strip(),
            str(event.get("MarginR", "")).strip(),
            str(event.get("MarginV", "")).strip(),
            str(event.get("Effect", "")).strip(),
            str(event.get("Text", "")).strip().replace("\n","\\N")                       # SRT style stored, ASS wanted
        ])
        ass_output.append(f"Dialogue: {event_line}")

    return "\n".join(ass_output)

def json_to_srt(json_data: dict, not_keep_index: bool = True) -> str:
    # print(json_data)
    events = json_data['Events']
    srt_output = []
    if not_keep_index:
        for i, event in tqdm(enumerate(events), desc="generating srt", total = len(events)):
            start_time = str(event.get("Start", "")).strip()  # Convert to SRT time format
            end_time = str(event.get("End", "")).strip()  # Convert to SRT time format

            start_time = timestamp_reform(start_time, output_format="srt")
            end_time = timestamp_reform(end_time, output_format="srt")
            text = event['Text']

            # Format the SRT entry
            srt_entry = f"{i + 1}\n{start_time} --> {end_time}\n{text}\n"
            srt_output.append(srt_entry)
    else:
        for event in tqdm(events, desc="generating srt", total = len(events)):
            index = str(event.get("Index", "")).strip()
            start_time = str(event.get("Start", "")).strip()  # Convert to SRT time format
            end_time = str(event.get("End", "")).strip()  # Convert to SRT time format

            start_time = timestamp_reform(start_time, output_format="srt")
            end_time = timestamp_reform(end_time, output_format="srt")
            text = event['Text']

            # Format the SRT entry
            srt_entry = f"{index}\n{start_time} --> {end_time}\n{text}\n"
            srt_output.append(srt_entry)

    return "\n".join(srt_output)

# Function to convert JSON to Word table format
def json_to_word(
        original_data: dict,
        translated_data: dict = None,
        reference_data: dict = None,
        force_disgard_index_and_time: bool = True
        ) -> Document:
    

    doc = Document()

    original_events = original_data['Events']

    # Create a table with 4 columns: Number, Time (start -> end), Subtitle Text, Empty
    table = doc.add_table(rows=len(original_events), cols=4 if reference_data is None else 5)
    table.style = 'Table Grid'

    # Loop through events and fill in the table
    for i, event in tqdm(enumerate(original_events), desc="filling original text in docx file", total=len(original_events)):
        index = event['Index']
        start_time = event['Start']
        end_time = event['End']
        subtitle_text = event['Text']

        row_cells = table.rows[i].cells

        row_cells[0].text = str(index)  # Number
        row_cells[1].text = f"{start_time} --> {end_time}"  # Time range
        row_cells[2].text = subtitle_text  # Subtitle text
        row_cells[3].text = ""  # Empty column

    if translated_data is not None:
        translated_events = translated_data['Events']
        # Loop through events and fill in the table
        for i, event in tqdm(enumerate(translated_events), desc="filling translated text in docx file", total=len(translated_events)):
            index = event['Index']
            start_time = event['Start']
            end_time = event['End']
            subtitle_text = event['Text']

            row_cells = table.rows[i].cells

            is_index_same = row_cells[0].text == str(index)  # Number
            is_timestamp_same = row_cells[1].text == f"{start_time} --> {end_time}"  # Time range

            if force_disgard_index_and_time:
                row_cells[3].text = subtitle_text  # Empty column
            elif is_index_same and is_timestamp_same:
                row_cells[3].text = subtitle_text  # Empty column
            else:
                print("index or time match failed.")
                print("original data:")
                print(row_cells[0].text, row_cells[1].text, row_cells[2].text)
                print("translated data:")
                print(str(index), f"{start_time} --> {end_time}", subtitle_text)



    if reference_data is not None:
        reference_events = reference_data['Events']
        # Loop through events and fill in the table
        for i, event in tqdm(enumerate(reference_events), desc="filling reference in docx file", total=len(reference_events)):
            index = event['Index']
            start_time = event['Start']
            end_time = event['End']
            subtitle_text = event['Text']

            row_cells = table.rows[i].cells

            is_index_same = row_cells[0].text == str(index)  # Number
            is_timestamp_same = row_cells[1].text == f"{start_time} --> {end_time}"  # Time range

            if force_disgard_index_and_time:
                row_cells[4].text = subtitle_text  # Empty column
            elif is_index_same and is_timestamp_same:
                row_cells[4].text = subtitle_text  # Empty column
            else:
                print("index or time match failed.")
                print("original data:")
                print(row_cells[0].text, row_cells[1].text, row_cells[2].text)
                print("translated data:")
                print(str(index), f"{start_time} --> {end_time}", subtitle_text)



    # set_font_based_on_characters(doc)
    doc = change_font(doc)

    return doc

# 函数: 修改字体
def change_font(doc: Document) -> Document:

    halfwidth_chars = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789!\"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~'

    # 处理段落
    for paragraph in tqdm(doc.paragraphs,desc="changing fonts of plain text"):
        for run in paragraph.runs:
            for ch in run.text:
                if ch in halfwidth_chars:
                    # 半角字符使用 Times New Roman
                    run.font.name = u'Times New Roman'
                    # # 设置中文部分的字体为宋体
                    # run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                else:
                    # 其他字符（假设为中文）使用 宋体
                    run.font.name = '宋体'
                    # 设置中文部分的字体为宋体
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 处理表格
    for table in doc.tables:
        for row in tqdm(table.rows, desc="changing fonts in tables"):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for ch in run.text:
                            if ch in halfwidth_chars:
                                # 半角字符使用 Times New Roman
                                run.font.name = u'Times New Roman'
                                # 设置中文部分的字体为宋体
                                # run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                            else:
                                # 其他字符（假设为中文）使用 宋体
                                run.font.name = '宋体'
                                # 设置中文部分的字体为宋体
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    return doc

def convert_subtitle_to_json(subtitle_file_path, json_file_path):
    if os.path.exists(subtitle_file_path):
        file_extension = os.path.splitext(subtitle_file_path)[1].lower()

    if file_extension == ".ass":
        data = parse_ass_file(subtitle_file_path)
    elif file_extension == ".srt":
        data = parse_srt_file(subtitle_file_path)
    elif file_extension == ".docx":
        data = parse_docx_file(subtitle_file_path)

    with open(json_file_path, 'w', encoding='utf-8') as json_file:
        json.dump(data, json_file, ensure_ascii=False, indent=4)

# Example usage
# ass_file = "example.ass"  # Path to your ASS file
# json_file = "X0643_KIRA the World EP02：少年们的 STAR MT 比赛_pt.json"  # Path where you want to save the JSON output
# convert_subtitle_to_json(ass_file, json_file)

# data = parse_ass_file("星光闪耀的少年_片花_ptX0861_EP05抢先看：少年们再跳主题曲 ！又哭又笑_简体中文_语料1.ass")
# word = json_to_word(data)
# word.save('sample_font.docx')
