import json
from docx import Document

# # Example JSON data (you can replace this with your actual JSON data)
# json_data = """
# {
#   "script_info": {
#     "title": "Example Subtitle",
#     "original_script": "OpenAI",
#     "script_type": "v4.00",
#     "collisions": "Normal"
#   },
#   "events": [
#     {
#       "start": "0:00:01.00",
#       "end": "0:00:04.00",
#       "style": "Default",
#       "text": "Welcome to the world of ASS subtitles!"
#     },
#     {
#       "start": "0:00:05.00",
#       "end": "0:00:08.00",
#       "style": "Default",
#       "text": "Here we are converting to JSON."
#     }
#   ],
#   "styles": [
#     {
#       "name": "Default",
#       "fontname": "Arial",
#       "fontsize": 20,
#       "primary_colour": "&H00FFFFFF",
#       "secondary_colour": "&H00000000",
#       "back_colour": "&H80000000",
#       "bold": -1,
#       "italic": 0,
#       "border_style": 1,
#       "outline": 1.0,
#       "shadow": 0.0,
#       "alignment": 2,
#       "margin_left": 10,
#       "margin_right": 10,
#       "margin_vertical": 10
#     }
#   ]
# }
# """


# from docx import Document

# # 创建一个新的 Word 文档
# doc = Document()

# # 添加一个有样式的表格
# table = doc.add_table(rows=3, cols=3)

# # 设置表格样式为 "Table Grid"（表格带网格线）
# table.style = 'Table Grid'

# # 添加数据
# data = [
#     ['Alice', 24, '济南'],
#     ['Bob', 27, 'Los Angeles'],
#     ['Charlie', 22, 'Chicago']
# ]

# for i, row_data in enumerate(data):
#     row_cells = table.rows[i].cells
#     for j, cell_data in enumerate(row_data):
#         row_cells[j].text = str(cell_data)

# # 保存文档
# doc.save('styled_table.docx')

# # Convert JSON string to Python dictionary
# data = json.loads(json_data)

# # Function to convert JSON to Word table format
# def convert_json_to_word(json_data: dict) -> Document:
#     events = json_data['Events']
#     doc = Document()
    
#     # Create a table with 4 columns: Number, Time (start -> end), Subtitle Text, Empty
#     table = doc.add_table(rows=len(events), cols=4)
#     table.style = 'Table Grid'

#     # Access the table's header row (we will not use it, just add a blank row)
#     # for row in table.rows:
#     #     for cell in row.cells:
#     #         cell.text = ""
    
#     # Loop through events and fill in the table
#     for i, event in enumerate(events):
#         start_time = event['Start']
#         end_time = event['End']
#         subtitle_text = event['Text']
        
#         # Add a new row
#         # row_cells = table.add_row().cells
#         row_cells = table.rows[i].cells
#         # Fill in the row with appropriate values
#         row_cells[0].text = str(i)  # Number
#         row_cells[1].text = f"{start_time} --> {end_time}"  # Time range
#         row_cells[2].text = subtitle_text  # Subtitle text
#         row_cells[3].text = ""  # Empty column
    
#     return doc
#     # Save the document
#     doc.save('subtitles.docx')
#     print("Word document 'subtitles.docx' has been created.")

from docx import Document
import re
from docx.shared import Pt

# Function to detect if text contains Chinese characters
def contains_chinese(text: str) -> bool:
    return bool(re.search(r'[\u4e00-\u9fff]', text))

# Function to apply different fonts based on text content
def set_font(run, text: str):
    if contains_chinese(text):
        run.font.name = '宋体'  # Chinese font
    else:
        run.font.name = 'Times New Roman'  # Western font

def convert_json_to_word(json_data: dict) -> Document:
    events = json_data['Events']
    doc = Document()

    # Create a table with 4 columns: Number, Time (start -> end), Subtitle Text, Empty
    table = doc.add_table(rows=len(events), cols=4)
    table.style = 'Table Grid'

    # Loop through events and fill in the table
    for i, event in enumerate(events):
        start_time = event['Start']
        end_time = event['End']
        subtitle_text = event['Text']
        
        row_cells = table.rows[i].cells

        # Fill Number column
        number_run = row_cells[0].paragraphs[0].add_run(str(i))  # Number
        set_font(number_run, str(i))  # Apply font style
        
        # Fill Time range column
        time_run = row_cells[1].paragraphs[0].add_run(f"{start_time} --> {end_time}")  # Time range
        set_font(time_run, f"{start_time} --> {end_time}")  # Apply font style
        
        # Fill Subtitle Text column
        subtitle_run = row_cells[2].paragraphs[0].add_run(subtitle_text)  # Subtitle text
        set_font(subtitle_run, subtitle_text)  # Apply font style
        
        # Empty column (No changes needed)
        row_cells[3].text = ""  # Empty column

    return doc

# Example usage
input_json_file = 'testfile\X0643_KIRA the World EP02：少年们的 STAR MT 比赛_pt.json'  # Path to the input JSON file
output_ass_file = 'output_test.docx'  # Path to save the output .ass file

# Read the JSON file
with open(input_json_file, 'r', encoding='utf-8') as f:
    json_data = json.load(f)
# Convert the JSON data to a Word document
doc = convert_json_to_word(json_data)

doc.save(output_ass_file)
