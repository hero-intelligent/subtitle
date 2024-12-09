from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def convert_json_to_word(json_data: dict) -> Document:
    events = json_data['Events']
    doc = Document()
    
    # Create a table with 4 columns: Number, Time (start -> end), Subtitle Text, Empty
    table = doc.add_table(rows=len(events) + 1, cols=4)  # +1 for header row
    table.style = 'Table Grid'
    
    # Set header row text
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Number'
    header_cells[1].text = 'Time (Start -> End)'
    header_cells[2].text = 'Subtitle Text'
    header_cells[3].text = 'Empty'
    
    # Loop through events and fill in the table
    for i, event in enumerate(events):
        start_time = event['Start']
        end_time = event['End']
        subtitle_text = event['Text']
        
        row_cells = table.rows[i + 1].cells  # Skip header row

        row_cells[0].text = str(i + 1)  # Number
        row_cells[1].text = f"{start_time} --> {end_time}"  # Time range
        row_cells[2].text = subtitle_text  # Subtitle text
        row_cells[3].text = ""  # Empty column
        
        # Apply different fonts for Chinese and English text
        # Set fonts for subtitle_text
        run = row_cells[2].paragraphs[0].runs[0]
        for char in subtitle_text:
            # Check if the character is Chinese (Unicode range for CJK characters)
            if '\u4e00' <= char <= '\u9fff':
                run.font.name = 'SimSun'  # Chinese font: SimSun (宋体)
                run._r.get_or_add_rPr().set(qn('w:eastAsia'), 'SimSun')
            else:
                run.font.name = 'Times New Roman'  # Western font: Times New Roman
            run.font.size = Pt(12)  # Font size

    # Adjust column widths (optional, depending on content)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Pt(100)  # You can adjust this as needed
    
    return doc
