from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 创建一个新的文档
doc = Document()

# 添加一段中文和英文混排的文本
p = doc.add_paragraph()

# 中文部分使用 "SimSun" 字体
run_chinese = p.add_run("这是中文部分。")
run_chinese.font.name = 'SimSun'  # 设置中文字体为宋体
# run_chinese.font.name = u'等线'  # 设置中文字体为宋体
run_chinese.font.size = Pt(12)

# 英文部分使用 "Times New Roman" 字体
run_english = p.add_run(" This is the English part.")
run_english.font.name = 'Times New Roman'  # 设置英文字体为 Times New Roman
run_english.font.size = Pt(12)

# 由于在 Word 中，中文字体和英文字体可能会有不同的显示方式，
# 所以我们需要使用 Microsoft Word 内部的 `w:eastAsia` 来设置中文字体
# 这样可以确保中文使用指定的字体。

# 获取到中文部分的字体设置对象，并设置字体
# run_chinese._r.get_or_add_rPr().set(qn('w:eastAsia'), u'等线')
run_chinese.element.rPr.rFonts.set(qn('w:eastAsia'), u'SimSun')

# 保存文档
doc.save('sample_font.docx')

print("文档已保存为 'sample_font.docx'")
