import pandas as pd
from docx import Document

# Your nested data structure
# data = {
#     "a": {"a": "a", "b": "b", "c": "c"},
#     "b": [
#         {"X": "X", "Y": "Y", "Z": "Z"},
#         {"X": "A", "Y": "B", "Z": "C"},
#         {"X": "D", "Y": "E", "Z": "F"},
#         {"X": "G", "Y": "H", "Z": "I"}
#     ]
# }
data = {
    'Name': ['Alice', 'Bob', 'Charlie'],
    'Age': [25, 30, 35],
    'City': ['New York', 'Los Angeles', 'Chicago']
}
# Step 1: Extract the list of dictionaries
b_data = data

# Step 2: Convert to a DataFrame
df_b = pd.DataFrame(b_data)

# Step 3: Create a Word document
doc = Document()
doc.add_heading('DataFrame Table from Nested Data', 0)

# Step 4: Add table to Word document
table = doc.add_table(rows=df_b.shape[0] + 1, cols=df_b.shape[1])

# Add headers to the first row
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(df_b.columns):
    hdr_cells[i].text = column_name

# Add data rows
for i, row in df_b.iterrows():
    row_cells = table.rows[i + 1].cells
    for j, value in enumerate(row):
        row_cells[j].text = str(value)

# Step 5: Save the document
doc.save('nested_data_table.docx')

print("Table has been added to the Word document.")
